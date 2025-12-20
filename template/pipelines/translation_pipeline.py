#!/usr/bin/env python3
"""
Machine Translation Pipeline using HuggingFace models
Supports both text and document translation with robust error handling
"""

import time
import torch
import numpy as np
from transformers import (
    AutoTokenizer, 
    AutoModelForSeq2SeqLM, 
    pipeline,
    MarianMTModel,
    MarianTokenizer
)
import io
import fitz  # PyMuPDF for PDF processing
from docx import Document
import tempfile
import os
from typing import Optional, Tuple, Dict, List, Union
import logging
from template.utils.hf_token import get_hf_token_dict

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class TranslationPipeline:
    """
    Machine translation pipeline using HuggingFace models.
    Supports multiple languages and document formats.
    """
    
    def __init__(self, model_name: str = "t5-small"):
        """
        Initialize the translation pipeline.
        
        Args:
            model_name: HuggingFace model name for translation
        """
        self.model_name = model_name
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        
        # Language code mapping for common languages
        self.language_codes = {
            "en": "english",
            "es": "spanish", 
            "fr": "french",
            "de": "german",
            "it": "italian",
            "pt": "portuguese",
            "ru": "russian",
            "ja": "japanese",
            "ko": "korean",
            "zh": "chinese",
            "ar": "arabic",
            "hi": "hindi",
            "nl": "dutch",
            "pl": "polish",
            "sv": "swedish",
            "tr": "turkish",
            "bg": "bulgarian",
            "ca": "catalan",
            "cs": "czech",
            "da": "danish",
            "el": "greek",
            "et": "estonian",
            "fi": "finnish",
            "hr": "croatian",
            "hu": "hungarian",
            "id": "indonesian",
            "lt": "lithuanian",
            "lv": "latvian",
            "ms": "malay",
            "mt": "maltese",
            "no": "norwegian",
            "ro": "romanian",
            "sk": "slovak",
            "sl": "slovenian",
            "th": "thai",
            "uk": "ukrainian",
            "vi": "vietnamese"
        }
        
        # Get HF token if available
        hf_token_kwargs = get_hf_token_dict()
        
        # Model loading with multiple fallback options
        fallback_models = [
            "t5-small",  # Simple, reliable T5 model
            "Helsinki-NLP/opus-mt-en-es",  # Marian model
            "facebook/mbart-large-50-many-to-many-mmt"  # Multilingual model
        ]
        
        for fallback_model in fallback_models:
            try:
                logger.info(f"🔄 Loading translation model: {fallback_model}")
                self.tokenizer = AutoTokenizer.from_pretrained(fallback_model, **hf_token_kwargs)
                self.model = AutoModelForSeq2SeqLM.from_pretrained(fallback_model, **hf_token_kwargs)
                self.model.to(self.device)
                self.model_name = fallback_model
                logger.info(f"✅ Translation model loaded successfully on {self.device}: {fallback_model}")
                
                # Try to use pipeline for easier inference
                try:
                    self.translation_pipeline = pipeline(
                        "translation", 
                        model=fallback_model, 
                        device=0 if self.device == "cuda" else -1,
                        token=hf_token_kwargs.get("token") if hf_token_kwargs else None
                    )
                    logger.info("✅ Translation pipeline initialized successfully")
                    break
                except Exception as e:
                    logger.warning(f"⚠️ Pipeline initialization failed, using manual inference: {e}")
                    self.translation_pipeline = None
                    break
                    
            except Exception as e:
                logger.warning(f"⚠️ Failed to load model {fallback_model}: {e}")
                continue
        else:
            # If all models failed, raise exception
            raise Exception(f"Could not load any translation model from: {fallback_models}")
        
        # Production settings
        self.max_chunk_size = 1000  # Maximum characters per translation chunk
        self.chunk_overlap = 50     # Characters overlap between chunks
        self.max_concurrent_chunks = 4  # Maximum concurrent chunk processing
        
        # Performance monitoring
        self.processing_stats = {
            'total_translations': 0,
            'total_text_length': 0,
            'total_processing_time': 0.0,
            'memory_usage_samples': []
        }
    
    def translate_text(
        self, 
        text: str, 
        source_language: str, 
        target_language: str,
        max_length: int = 512
    ) -> Tuple[str, float]:
        """
        Translate text from source language to target language.
        
        Args:
            text: Input text to translate
            source_language: Source language code
            target_language: Target language code
            max_length: Maximum length for translation chunks
            
        Returns:
            Tuple of (translated_text, processing_time)
        """
        start_time = time.time()
        
        try:
            logger.info(f"🌐 Translating text from {source_language} to {target_language}")
            logger.info(f"   Text length: {len(text)} characters")
            
            # Handle long texts by chunking
            if len(text) > max_length:
                logger.info(f"📝 Text is long ({len(text)} chars), chunking for translation")
                translated_chunks = []
                
                # Split text into chunks
                chunks = self._chunk_text(text, max_length)
                logger.info(f"   Split into {len(chunks)} chunks")
                
                for i, chunk in enumerate(chunks):
                    logger.info(f"   Translating chunk {i+1}/{len(chunks)} ({len(chunk)} chars)")
                    translated_chunk = self._translate_chunk(chunk, source_language, target_language)
                    translated_chunks.append(translated_chunk)
                
                # Combine translated chunks
                translated_text = " ".join(translated_chunks)
                logger.info(f"✅ Combined {len(translated_chunks)} translated chunks")
                
            else:
                # Direct translation for short texts
                translated_text = self._translate_chunk(text, source_language, target_language)
            
            processing_time = time.time() - start_time
            logger.info(f"✅ Translation completed in {processing_time:.2f}s")
            logger.info(f"   Original length: {len(text)} characters")
            logger.info(f"   Translated length: {len(translated_text)} characters")
            
            return translated_text, processing_time
            
        except Exception as e:
            processing_time = time.time() - start_time
            logger.error(f"❌ Translation failed: {e}")
            raise Exception(f"Translation failed: {str(e)}")
    
    def _translate_chunk(
        self, 
        text: str, 
        source_language: str, 
        target_language: str
    ) -> str:
        """Translate a single chunk of text"""
        try:
            if self.translation_pipeline:
                # Use pipeline if available
                result = self.translation_pipeline(text)
                return result[0]['translation_text']
            else:
                # Manual inference with T5-style prefix
                if "t5" in self.model_name.lower():
                    # T5 models need translation prefix
                    prefix = f"translate English to {target_language}: "
                    input_text = prefix + text
                elif "mbart" in self.model_name.lower():
                    # mBART models need language codes
                    input_text = text
                else:
                    # Marian models work directly
                    input_text = text
                
                inputs = self.tokenizer(
                    input_text, 
                    return_tensors="pt", 
                    max_length=512, 
                    truncation=True,
                    padding=True
                ).to(self.device)
                
                outputs = self.model.generate(
                    **inputs,
                    max_length=512,
                    num_beams=4,
                    early_stopping=True,
                    pad_token_id=self.tokenizer.eos_token_id
                )
                
                translated_text = self.tokenizer.decode(outputs[0], skip_special_tokens=True)
                return translated_text
                
        except Exception as e:
            logger.error(f"❌ Chunk translation failed: {e}")
            raise
    
    def _chunk_text(self, text: str, max_length: int) -> List[str]:
        """Split text into chunks for translation"""
        words = text.split()
        chunks = []
        current_chunk = []
        current_length = 0
        
        for word in words:
            word_length = len(word) + 1  # +1 for space
            if current_length + word_length > max_length and current_chunk:
                chunks.append(" ".join(current_chunk))
                current_chunk = [word]
                current_length = word_length
            else:
                current_chunk.append(word)
                current_length += word_length
        
        if current_chunk:
            chunks.append(" ".join(current_chunk))
        
        return chunks
    
    def translate_document(
        self, 
        file_data, 
        filename: str,
        source_language: str, 
        target_language: str
    ) -> Tuple[str, float, Dict]:
        """
        Translate document content from source language to target language.
        
        Args:
            file_data: Raw file bytes
            filename: Original filename
            source_language: Source language code
            target_language: Target language code
            
        Returns:
            Tuple of (translated_text, processing_time, metadata)
        """
        start_time = time.time()
        
        try:
            logger.info(f"📄 Translating document: {filename}")
            logger.info(f"   File size: {len(file_data)} bytes")
            logger.info(f"   From {source_language} to {target_language}")
            
            # Extract text from document based on file type
            file_extension = filename.lower().split('.')[-1]
            
            if file_extension == 'pdf':
                extracted_text = self._extract_text_from_pdf(file_data)
            elif file_extension == 'docx':
                extracted_text = self._extract_text_from_docx(file_data)
            elif file_extension in ['txt', 'text']:
                extracted_text = self._extract_text_from_txt(file_data)
            else:
                raise Exception(f"Unsupported file format: {file_extension}")
            
            logger.info(f"✅ Text extracted: {len(extracted_text)} characters")
            
            # Translate the extracted text
            translated_text, translation_time = self.translate_text(
                extracted_text, source_language, target_language
            )
            
            total_time = time.time() - start_time
            
            # Prepare metadata
            metadata = {
                'original_filename': filename,
                'file_extension': file_extension,
                'file_size_bytes': len(file_data),
                'extracted_text_length': len(extracted_text),
                'translated_text_length': len(translated_text),
                'source_language': source_language,
                'target_language': target_language,
                'extraction_time': total_time - translation_time,
                'translation_time': translation_time,
                'total_time': total_time
            }
            
            logger.info(f"✅ Document translation completed in {total_time:.2f}s")
            return translated_text, total_time, metadata
            
        except Exception as e:
            total_time = time.time() - start_time
            logger.error(f"❌ Document translation failed: {e}")
            raise Exception(f"Document translation failed: {str(e)}")
    
    def _extract_text_from_pdf(self, file_data: bytes) -> str:
        """Extract text from PDF file"""
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                temp_file.write(file_data)
                temp_file_path = temp_file.name
            
            try:
                # Open PDF and extract text
                doc = fitz.open(temp_file_path)
                text_parts = []
                
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    text = page.get_text()
                    if text.strip():
                        text_parts.append(text.strip())
                
                doc.close()
                
                extracted_text = "\n\n".join(text_parts)
                logger.info(f"✅ PDF text extraction: {len(doc)} pages, {len(extracted_text)} chars")
                return extracted_text
                
            finally:
                # Clean up temporary file
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                    
        except Exception as e:
            logger.error(f"❌ PDF text extraction failed: {e}")
            raise Exception(f"PDF text extraction failed: {str(e)}")
    
    def _extract_text_from_docx(self, file_data: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(file_data)
                temp_file_path = temp_file.name
            
            try:
                # Open DOCX and extract text
                doc = Document(temp_file_path)
                text_parts = []
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text.strip())
                
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_parts.append(cell.text.strip())
                
                extracted_text = "\n\n".join(text_parts)
                logger.info(f"✅ DOCX text extraction: {len(text_parts)} paragraphs, {len(extracted_text)} chars")
                return extracted_text
                
            finally:
                # Clean up temporary file
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                    
        except Exception as e:
            logger.error(f"❌ DOCX text extraction failed: {e}")
            raise Exception(f"DOCX text extraction failed: {str(e)}")
    
    def _extract_text_from_txt(self, file_data) -> str:
        """Extract text from text file"""
        try:
            # Handle both string and bytes input
            if isinstance(file_data, str):
                # If it's already a string, return it directly
                logger.info(f"✅ Text file already string: {len(file_data)} chars")
                return file_data
            elif isinstance(file_data, bytes):
                # If it's bytes, decode it
                encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
                
                for encoding in encodings:
                    try:
                        text = file_data.decode(encoding)
                        logger.info(f"✅ Text file decoded using {encoding}: {len(text)} chars")
                        return text
                    except UnicodeDecodeError:
                        continue
                
                # If all encodings fail, use latin-1 with replacement
                text = file_data.decode('latin-1', errors='replace')
                logger.warning(f"⚠️ Text file decoded with replacement: {len(text)} chars")
                return text
            else:
                # Convert other types to string
                text = str(file_data)
                logger.info(f"✅ Text file converted to string: {len(text)} chars")
                return text
                
        except Exception as e:
            logger.error(f"❌ Text file extraction failed: {e}")
            raise Exception(f"Text file extraction failed: {str(e)}")
    
    def get_supported_languages(self) -> List[str]:
        """Get list of supported language codes"""
        return list(self.language_codes.keys())
    
    def validate_language_pair(self, source_lang: str, target_lang: str) -> bool:
        """Validate if language pair is supported"""
        return source_lang in self.language_codes and target_lang in self.language_codes
    
    def get_model_info(self) -> Dict:
        """Get information about the loaded model"""
        return {
            'model_name': self.model_name,
            'device': self.device,
            'supported_languages': self.get_supported_languages(),
            'model_loaded': self.model is not None,
            'pipeline_available': self.translation_pipeline is not None
        }
    
    def _update_stats(self, text_length: int, processing_time: float):
        """Update processing statistics"""
        self.processing_stats['total_translations'] += 1
        self.processing_stats['total_text_length'] += text_length
        self.processing_stats['total_processing_time'] += processing_time
    
    def _cleanup_memory(self):
        """Clean up memory and perform garbage collection"""
        try:
            # Clear CUDA cache if using GPU
            if self.device == "cuda" and torch.cuda.is_available():
                torch.cuda.empty_cache()
            
            # Force garbage collection
            gc.collect()
            
            # Log memory usage
            process = psutil.Process()
            current_memory = process.memory_info().rss / (1024 * 1024 * 1024)
            self.processing_stats['memory_usage_samples'].append(current_memory)
            logger.info(f"🧹 Memory cleanup completed. Current usage: {current_memory:.2f}GB")
            
        except Exception as e:
            logger.warning(f"⚠️ Memory cleanup failed: {e}")
    
    def get_performance_stats(self) -> Dict:
        """Get performance statistics"""
        return {
            'total_translations': self.processing_stats['total_translations'],
            'total_text_length': self.processing_stats['total_text_length'],
            'total_processing_time': self.processing_stats['total_processing_time'],
            'average_processing_speed': (
                self.processing_stats['total_text_length'] / 
                max(self.processing_stats['total_processing_time'], 0.001)
            ),
            'memory_usage_samples': self.processing_stats['memory_usage_samples'][-10:],  # Last 10 samples
            'model_info': {
                'name': self.model_name,
                'device': self.device,
                'max_chunk_size': self.max_chunk_size,
                'chunk_overlap': self.chunk_overlap,
                'max_concurrent_chunks': self.max_concurrent_chunks
            }
        }
    
    def optimize_for_production(self, target_max_chunk_size: int = 800, target_chunk_overlap: int = 30):
        """Optimize pipeline settings for production use"""
        self.max_chunk_size = target_max_chunk_size
        self.chunk_overlap = target_chunk_overlap
        
        logger.info(f"⚙️ Translation pipeline optimized for production:")
        logger.info(f"   Max chunk size: {target_max_chunk_size} characters")
        logger.info(f"   Chunk overlap: {target_chunk_overlap} characters")
        
        # Force memory cleanup
        self._cleanup_memory()

# Global instance
# Global instance for backward compatibility (lazy initialization)
# Removed immediate instantiation to prevent model loading during import
_translation_pipeline_instance = None

def get_translation_pipeline_instance():
    """Get or create the global translation pipeline instance (lazy)"""
    global _translation_pipeline_instance
    if _translation_pipeline_instance is None:
        _translation_pipeline_instance = TranslationPipeline()
    return _translation_pipeline_instance

# For backward compatibility - but prefer creating new instances
translation_pipeline = None
